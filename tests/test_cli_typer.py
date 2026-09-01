"""Tests for CLI with Typer - tailor and re-tailor commands."""

from datetime import date
from unittest.mock import AsyncMock, MagicMock, patch
import os

import pytest
from typer.testing import CliRunner

from sira.main import app
from sira.models.agents.output import CV, WorkExperience, ScrapedJobPosting
from sira.models.workflow import ResumeTailorResult

runner = CliRunner()


def _make_cv(full_name: str = "Jane Doe") -> CV:
    return CV(
        full_name=full_name,
        contact_info="jane@example.com",
        summary="Platform engineer.",
        skills=["Python", "SQL"],
        experience=[
            WorkExperience(
                company="Acme",
                role="Engineer",
                dates="2022-2026",
                highlights=["Built services"],
            )
        ],
        education=["BSc CS"],
    )


def _make_result(cv: CV | None = None, passed: bool = True) -> ResumeTailorResult:
    cv = cv or _make_cv()
    from sira.models.agents.output import CVDiff, FinalReport, GapAnalysis

    final_report = FinalReport(
        company_name="Acme Corp",
        job_title="Software Engineer",
        generated_at="2026-01-01T00:00:00Z",
        overall_recommendation="Strong Match",
        match_score=85,
        what_changed=CVDiff(sections_modified=["summary"]),
        gaps=GapAnalysis(
            covered_keywords=["Python"],
            missing_keywords=["Rust"],
        ),
        suggestions_to_strengthen=["Add Rust experience"],
        audit_summary="Looks good.",
        recommendation_rationale="Strong alignment with role requirements.",
        passed=passed,
    )
    return ResumeTailorResult(
        company_name="Acme Corp",
        job_title="Software Engineer",
        tailored_resume=cv.model_dump_json(),
        audit_report={
            "passed": passed,
            "hallucination_score": 0,
            "ai_cliche_score": 1,
            "issues": [],
            "feedback_summary": "Looks good.",
        },
        passed=passed,
        final_report=final_report,
    )


def _make_scraped_job(
    markdown: str = "# Job Posting\nPython engineer",
) -> ScrapedJobPosting:
    return ScrapedJobPosting(
        markdown=markdown,
        url="https://example.com/job/123",
        source_text="Raw job posting text",
        extraction_strategy="test",
    )


def test_tailor_command_success(tmp_path, monkeypatch) -> None:
    """tailor command with valid inputs should succeed and show output."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(
        return_value=str(output_dir / "tailored_resume_acme_corp.md")
    )

    scraped_job = _make_scraped_job()

    with (
        patch(
            "sira.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        # mock_parser is no longer used since mock_parser_cls is now _
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    assert "✅ Job completed" in result.output
    mock_workflow.run.assert_called_once()
    mock_generate_resume.assert_called_once()


def test_tailor_command_invalid_url_format(tmp_path, monkeypatch) -> None:
    """tailor command with invalid URL format should return 1."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    result = runner.invoke(
        app,
        [
            "tailor",
            "not-a-valid-url",
            str(resume_file),
            "--output-dir",
            str(output_dir),
        ],
    )

    assert result.exit_code == 1
    assert "http://" in result.output or "https://" in result.output


def test_tailor_command_resume_not_found(tmp_path, monkeypatch) -> None:
    """tailor command with non-existent resume should return 1."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    result = runner.invoke(
        app,
        [
            "tailor",
            "https://example.com/job/123",
            "/no/such/file.md",
            "--output-dir",
            str(output_dir),
        ],
    )

    assert result.exit_code == 1
    assert "not found" in result.output.lower() or "❌" in result.output


def test_tailor_command_scraping_failure(tmp_path, monkeypatch) -> None:
    """tailor command when scraping fails should return 1."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    mock_scraper = AsyncMock(side_effect=Exception("Network error"))

    with patch("sira.main.job_scraper_agent.run", mock_scraper):
        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "scrape" in result.output.lower() or "failed" in result.output.lower()


def test_tailor_command_failed_audit_exits_zero(tmp_path, monkeypatch) -> None:
    """tailor command with failed audit should exit 0 (report still generated)."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=False)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(
        return_value=str(output_dir / "tailored_resume_acme_corp.md")
    )

    scraped_job = _make_scraped_job()

    with (
        patch(
            "sira.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        # mock_parser is no longer used since mock_parser_cls is now _
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    mock_workflow.run.assert_called_once()


def test_tailor_command_empty_job_content(tmp_path, monkeypatch) -> None:
    """tailor command when scraped job is empty should return 1."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    scraped_job = _make_scraped_job(markdown="")

    with patch(
        "sira.main.job_scraper_agent.run",
        AsyncMock(return_value=MagicMock(output=scraped_job)),
    ):
        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "empty" in result.output.lower()


def test_tailor_command_docx_conversion(tmp_path, monkeypatch) -> None:
    """tailor command should convert DOCX resume to markdown."""
    # Use docx library to create a minimal docx file.
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    resume_file = tmp_path / "resume.docx"
    doc = Document()
    doc.add_heading("Jane Doe", 0)
    doc.add_paragraph("Python developer")
    doc.save(str(resume_file))

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(
        return_value=str(output_dir / "tailored_resume_acme_corp.md")
    )

    scraped_job = _make_scraped_job()

    with (
        patch(
            "sira.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        # mock_parser is no longer used since mock_parser_cls is now _
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    mock_workflow.run.assert_called_once()
    converted_file = output_dir / "resume_converted.md"
    assert converted_file.exists()


# --- re_tailor tests ---


def test_re_tailor_success(tmp_path, monkeypatch) -> None:
    """re_tailor with valid job ID and recommendations should succeed."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(
        return_value=str(output_dir / "tailored_resume_acme_corp.md")
    )

    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Smith\n\nPython developer", encoding="utf-8")

    with (
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting\nPython engineer",
        )
        mock_repo.get_source_by_id.return_value = MagicMock(path=str(resume_file))
        mock_repo_cls.return_value = mock_repo

        # mock_parser is no longer used since mock_parser_cls is now _
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123", path=str(resume_file)),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add more detail about leadership skills",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    assert "✅ Re-tailoring completed" in result.output
    mock_workflow.run.assert_called_once()
    mock_repo.save_tailored_resume.assert_called_once()


def test_re_tailor_job_not_found(tmp_path, monkeypatch) -> None:
    """re_tailor with non-existent job ID should return 1."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    with patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = None
        mock_repo_cls.return_value = mock_repo

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "nonexistent-id",
                "Add leadership skills",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "Job not found" in result.output


def test_re_tailor_no_job_markdown(tmp_path, monkeypatch) -> None:
    """re_tailor should fail when stored job has no posting markdown."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\n\nPython developer.", encoding="utf-8")

    with patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="",
        )
        mock_repo.get_source_by_id.return_value = MagicMock(path=str(resume_file))
        mock_repo_cls.return_value = mock_repo

        with (
            patch("sira.main.PydanticAIResumeParser") as _,
            patch("sira.main.ResumeMemoryService") as mock_svc_cls,
        ):
            mock_svc = MagicMock()
            mock_svc.resolve_original_resume.return_value = MagicMock(
                source=MagicMock(id="src-123", path=str(resume_file)),
                cv=_make_cv(),
            )
            mock_svc.aresolve_original_resume = AsyncMock(
                return_value=mock_svc.resolve_original_resume.return_value
            )
            mock_svc_cls.return_value = mock_svc

            result = runner.invoke(
                app,
                [
                    "re-tailor",
                    "job-456",
                    "Add leadership skills",
                    "--output-dir",
                    str(output_dir),
                ],
            )

    assert result.exit_code == 1
    assert "No job posting content stored" in result.output


def test_re_tailor_with_resume_path(tmp_path, monkeypatch) -> None:
    """re_tailor should use explicit resume path when provided."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(
        return_value=str(output_dir / "tailored_resume_acme_corp.md")
    )

    with (
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting\nPython engineer",
        )
        mock_repo_cls.return_value = mock_repo

        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add leadership skills",
                "--resume-path",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    # When --resume-path is passed, we should not call get_source_by_id.
    mock_repo.get_source_by_id.assert_not_called()
    mock_svc.aresolve_original_resume.assert_awaited_with(path=str(resume_file))


def test_re_tailor_missing_resume_file(tmp_path, monkeypatch) -> None:
    """re_tailor should fail when explicit resume path does not exist."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    with patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting",
        )
        mock_repo_cls.return_value = mock_repo

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add skills",
                "--resume-path",
                "/no/such/file.md",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "not found" in result.output.lower() or "❌" in result.output


def test_re_tailor_missing_original_source_file(tmp_path, monkeypatch) -> None:
    """re_tailor should fail when original source file is gone and no --resume-path provided."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    with patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting",
        )
        # Source exists in DB but file is gone
        mock_repo.get_source_by_id.return_value = MagicMock(
            path="/nonexistent/resume.md"
        )
        mock_repo_cls.return_value = mock_repo

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add skills",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "Original resume not found at recorded path" in result.output
    assert "--resume-path" in result.output


def test_tailor_command_custom_patterns(tmp_path, monkeypatch) -> None:
    """tailor command with custom patterns should create correctly named files."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)

    captured_args = {}

    def mock_generate_resume(result, output_dir, base_filename):
        captured_args["output_dir"] = output_dir
        captured_args["base_filename"] = base_filename
        return os.path.join(output_dir, f"{base_filename}.md")

    scraped_job = _make_scraped_job()

    with (
        patch(
            "sira.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        # mock_parser is no longer used since mock_parser_cls is now _
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
                "--output-pattern",
                "{company_name}-{timestamp}",
                "--resume-name-pattern",
                "{full_name}-{job_title}",
            ],
        )

    assert result.exit_code == 0, result.output
    mock_workflow.run.assert_called_once()

    today = date.today().strftime("%Y%m%d")
    expected_dir = str(output_dir / f"acme_corp-{today}")
    expected_base = "jane_doe-software_engineer"
    assert captured_args["output_dir"] == expected_dir
    assert captured_args["base_filename"] == expected_base


def test_tailor_accepts_fast_and_gate_flags():
    """New CLI flags are recognized (no error from Typer parsing)."""
    import re

    from typer.testing import CliRunner
    from sira.main import app

    runner = CliRunner()
    result = runner.invoke(app, ["tailor", "--help"])
    assert result.exit_code == 0
    # Strip ANSI styling before matching: in color mode (e.g. CI, where rich
    # emits color) each option's leading "--" is split by a style reset
    # ("-\x1b[0m\x1b[1;36m-fast"), so the literal "--fast" substring is absent
    # until the escape codes are removed. Locally output is uncolored and passes
    # either way, which is why this only failed in CI.
    output = re.sub(r"\x1b\[[0-9;]*m", "", result.output)
    assert "--fast" in output
    assert "--write-attempts" in output
    assert "--review-iterations" in output
    assert "--no-quality-gate" in output or "--quality-gate" in output
    assert "--gate-threshold" in output


def test_re_tailor_custom_patterns(tmp_path, monkeypatch) -> None:
    """re_tailor with custom patterns should create correctly named files."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)

    captured_args = {}

    def mock_generate_resume(result, output_dir, base_filename):
        captured_args["output_dir"] = output_dir
        captured_args["base_filename"] = base_filename
        return os.path.join(output_dir, f"{base_filename}.md")

    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Smith\n\nPython developer", encoding="utf-8")

    with (
        patch("sira.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("sira.main.generate_resume", mock_generate_resume),
        patch("sira.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("sira.main.PydanticAIResumeParser") as _,
        patch("sira.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting\nPython engineer",
        )
        mock_repo.get_source_by_id.return_value = MagicMock(path=str(resume_file))
        mock_repo_cls.return_value = mock_repo

        # mock_parser is no longer used since mock_parser_cls is now _
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123", path=str(resume_file)),
            cv=cv,
        )
        mock_svc.aresolve_original_resume = AsyncMock(
            return_value=mock_svc.resolve_original_resume.return_value
        )
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add more detail about leadership skills",
                "--output-dir",
                str(output_dir),
                "--output-pattern",
                "{timestamp}-{company_name}",
                "--resume-name-pattern",
                "{job_title}-{full_name}",
            ],
        )

    assert result.exit_code == 0, result.output
    mock_workflow.run.assert_called_once()

    today = date.today().strftime("%Y%m%d")
    expected_dir = str(output_dir / f"{today}-acme_corp")
    expected_base = "software_engineer-jane_doe"
    assert captured_args["output_dir"] == expected_dir
    assert captured_args["base_filename"] == expected_base
